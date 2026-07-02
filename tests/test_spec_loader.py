import tempfile
import unittest
from pathlib import Path

from spec_loader import load_spec_as_markdown


TEXT_CASES = [
    (
        "standard.md",
        "# AXI Bridge Spec\n\n## 1. Overview\nAXI to SRAM bridge.\n\n## 2. Interface\nAWVALID is supported.\n",
        ["# AXI Bridge Spec", "## 1. Overview", "## 2. Interface"],
    ),
    (
        "numbered.txt",
        "AXI Bridge Specification\n\n1. Module Overview\nThe bridge converts AXI to SRAM.\n\n1.1 Interface Signals\nAWVALID and WVALID are supported.\n",
        ["# AXI Bridge Specification", "# 1 Module Overview", "## 1.1 Interface Signals"],
    ),
    (
        "nested_numbered.txt",
        "SRAM Controller Spec\n\n1 Architecture\nText.\n1.1 Write Path\nText.\n1.1.1 Write Response\nText.\n2 Reset\nText.\n",
        ["# SRAM Controller Spec", "# 1 Architecture", "## 1.1 Write Path", "### 1.1.1 Write Response"],
    ),
    (
        "chinese_chapter.txt",
        "DMA Bridge 设计规范\n\n第一章 模块概述\n描述模块用途。\n第二章 接口信号\n描述 AXI 接口。\n",
        ["# DMA Bridge 设计规范", "# 第一章 模块概述", "# 第二章 接口信号"],
    ),
    (
        "chinese_number.txt",
        "GPIO 模块规范\n\n一、模块概述\n支持寄存器访问。\n二、异常处理\n非法地址返回错误。\n",
        ["# GPIO 模块规范", "# 一、模块概述", "# 二、异常处理"],
    ),
    (
        "setext.txt",
        "Timer Module Specification\n==========================\n\nRegister Map\n------------\nCTRL bit0 enables timer.\n",
        ["# Timer Module Specification", "## Register Map"],
    ),
    (
        "all_caps.txt",
        "UART MODULE SPEC\n\nINTERFACE SIGNALS\nclk and rstn are defined.\nRESET BEHAVIOR\nreset clears fifo.\n",
        ["## UART MODULE SPEC", "## INTERFACE SIGNALS", "## RESET BEHAVIOR"],
    ),
    (
        "bullets.txt",
        "Interrupt Controller Spec\n\n1. Overview\n- 1. this bullet should stay a bullet\n2. Registers\nIRQ_STATUS reports pending interrupts.\n",
        ["# Interrupt Controller Spec", "# 1 Overview", "- 1. this bullet should stay a bullet", "# 2 Registers"],
    ),
    (
        "table_text.txt",
        "Register Block Spec\n\n1. Register Map\n| name | offset |\n| --- | --- |\n| CTRL | 0x0 |\n",
        ["# Register Block Spec", "# 1 Register Map", "| CTRL | 0x0 |"],
    ),
    (
        "unstructured.txt",
        "This block describes reset behavior. aresetn low clears all valid signals.\n",
        ["## Content", "aresetn low clears all valid signals"],
    ),
]


class SpecLoaderTextTests(unittest.TestCase):
    def test_text_and_markdown_cases(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            self.assertGreaterEqual(len(TEXT_CASES), 10)
            for filename, content, expected_fragments in TEXT_CASES:
                with self.subTest(filename=filename):
                    path = root / filename
                    path.write_text(content, encoding="utf-8")

                    loaded = load_spec_as_markdown(path)

                    for fragment in expected_fragments:
                        self.assertIn(fragment, loaded.markdown_text)
                    self.assertGreaterEqual(loaded.report.heading_count, 1)

    def test_standard_markdown_is_not_rewritten(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "standard.md"
            content = "# Title\n\n## Child\nBody\n"
            path.write_text(content, encoding="utf-8")

            loaded = load_spec_as_markdown(path)

            self.assertEqual(content, loaded.markdown_text)
            self.assertEqual(loaded.report.source_format, "markdown")
            self.assertEqual(loaded.report.inferred_heading_count, 0)

    def test_unsupported_suffix_fails_clearly(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "spec.rtf"
            path.write_text("1. Overview\nText", encoding="utf-8")

            with self.assertRaisesRegex(ValueError, "不支持的输入格式"):
                load_spec_as_markdown(path)


class SpecLoaderDocxTests(unittest.TestCase):
    def test_docx_heading_paragraphs_and_table(self):
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx is not installed")

        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "bridge.docx"
            document = Document()
            document.add_heading("AXI Bridge Specification", level=1)
            document.add_heading("1. Interface Signals", level=2)
            document.add_paragraph("AWVALID and WVALID are supported.")
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Signal"
            table.cell(0, 1).text = "Direction"
            table.cell(1, 0).text = "AWVALID"
            table.cell(1, 1).text = "input"
            document.save(str(path))

            loaded = load_spec_as_markdown(path)

            self.assertIn("# AXI Bridge Specification", loaded.markdown_text)
            self.assertIn("## 1. Interface Signals", loaded.markdown_text)
            self.assertIn("| Signal | Direction |", loaded.markdown_text)
            self.assertEqual(loaded.report.source_format, "docx")
            self.assertEqual(loaded.report.table_count, 1)

    def test_docx_numbered_plain_paragraphs(self):
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx is not installed")

        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "plain.docx"
            document = Document()
            document.add_paragraph("Timer Module Specification")
            document.add_paragraph("1. Overview")
            document.add_paragraph("Timer has enable and interrupt output.")
            document.add_paragraph("1.1 Reset")
            document.add_paragraph("reset clears CTRL register.")
            document.save(str(path))

            loaded = load_spec_as_markdown(path)

            self.assertIn("# Timer Module Specification", loaded.markdown_text)
            self.assertIn("# 1 Overview", loaded.markdown_text)
            self.assertIn("## 1.1 Reset", loaded.markdown_text)


class SpecLoaderPdfTests(unittest.TestCase):
    def test_pdf_numbered_text(self):
        try:
            import fitz
        except ImportError:
            self.skipTest("PyMuPDF is not installed")

        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "pdf_numbered.pdf"
            document = fitz.open()
            page = document.new_page()
            page.insert_text(
                (72, 72),
                "AXI Bridge Specification\n\n1. Overview\nAXI to SRAM bridge.\n\n1.1 Reset\nreset clears valid signals.",
            )
            document.save(str(path))
            document.close()

            loaded = load_spec_as_markdown(path)

            self.assertEqual(loaded.report.source_format, "pdf")
            self.assertIn("# AXI Bridge Specification", loaded.markdown_text)
            self.assertIn("# 1 Overview", loaded.markdown_text)
            self.assertIn("## 1.1 Reset", loaded.markdown_text)

    def test_pdf_without_headings_gets_default_section(self):
        try:
            import fitz
        except ImportError:
            self.skipTest("PyMuPDF is not installed")

        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "flat.pdf"
            document = fitz.open()
            page = document.new_page()
            page.insert_text((72, 72), "aresetn low clears all AXI valid signals.")
            document.save(str(path))
            document.close()

            loaded = load_spec_as_markdown(path)

            self.assertIn("## Content", loaded.markdown_text)
            self.assertTrue(loaded.report.warnings)


if __name__ == "__main__":
    unittest.main()
