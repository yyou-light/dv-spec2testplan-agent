import re
from dataclasses import dataclass, field
from pathlib import Path


SUPPORTED_SUFFIXES = {".md", ".markdown", ".txt", ".text", ".docx", ".pdf"}


@dataclass
class SpecLoadReport:
    source_path: str
    source_format: str
    heading_count: int
    inferred_heading_count: int = 0
    table_count: int = 0
    warnings: list[str] = field(default_factory=list)

    @property
    def confidence(self) -> str:
        if self.heading_count >= 3 and not self.warnings:
            return "high"
        if self.heading_count >= 1:
            return "medium"
        return "low"


@dataclass
class LoadedSpec:
    markdown_text: str
    report: SpecLoadReport


def load_spec_as_markdown(path: str | Path) -> LoadedSpec:
    source_path = Path(path)
    suffix = source_path.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ValueError(
            f"不支持的输入格式: {suffix or '<无扩展名>'}。"
            "当前支持 .md/.markdown/.txt/.text/.docx/.pdf。"
        )

    if suffix in {".md", ".markdown"}:
        text = read_text_file(source_path)
        return normalize_text_to_markdown(text, source_path, source_format="markdown")

    if suffix in {".txt", ".text"}:
        text = read_text_file(source_path)
        return normalize_text_to_markdown(text, source_path, source_format="text")

    if suffix == ".docx":
        text, table_count = read_docx_as_markdown(source_path)
        loaded = normalize_text_to_markdown(text, source_path, source_format="docx")
        loaded.report.table_count = table_count
        return loaded

    text = read_pdf_text(source_path)
    return normalize_text_to_markdown(text, source_path, source_format="pdf")


def read_text_file(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise UnicodeDecodeError("unknown", b"", 0, 1, f"无法识别文本编码: {path}")


def read_docx_as_markdown(path: Path) -> tuple[str, int]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("读取 .docx 需要安装 python-docx: pip install python-docx") from exc

    document = Document(str(path))
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (paragraph.style.name or "").lower()
        heading_match = re.search(r"heading\s+(\d+)", style_name)
        if heading_match:
            level = min(max(int(heading_match.group(1)), 1), 6)
            lines.append(f"{'#' * level} {text}")
        else:
            lines.append(text)

    table_count = 0
    for table in document.tables:
        table_count += 1
        rows = [[cell.text.strip().replace("\n", " ") for cell in row.cells] for row in table.rows]
        lines.extend(markdown_table(rows))

    return "\n\n".join(lines), table_count


def read_pdf_text(path: Path) -> str:
    try:
        import fitz
    except ImportError as exc:
        raise RuntimeError("读取 .pdf 需要安装 PyMuPDF: pip install PyMuPDF") from exc

    chunks: list[str] = []
    with fitz.open(str(path)) as document:
        for page_index, page in enumerate(document, start=1):
            text = page.get_text("text").strip()
            if text:
                chunks.append(f"\n\n<!-- page {page_index} -->\n\n{text}")
    if not chunks:
        raise ValueError(f"PDF 未提取到可用文本: {path}")
    return "\n".join(chunks)


def normalize_text_to_markdown(text: str, source_path: Path, source_format: str) -> LoadedSpec:
    text = normalize_line_endings(text)
    report = SpecLoadReport(
        source_path=str(source_path),
        source_format=source_format,
        heading_count=count_markdown_headings(text),
    )

    if report.heading_count:
        markdown_text = ensure_single_trailing_newline(text)
        return LoadedSpec(markdown_text=markdown_text, report=report)

    markdown_text, inferred = infer_markdown_headings(text)
    report.inferred_heading_count = inferred
    report.heading_count = count_markdown_headings(markdown_text)

    if report.heading_count == 0:
        report.warnings.append("未识别到章节标题，已将全文放入默认章节。")
        title = source_path.stem.replace("_", " ").replace("-", " ").strip() or "Unstructured Spec"
        markdown_text = f"# {title}\n\n## Content\n\n{text.strip()}\n"
        report.heading_count = 2
    elif inferred:
        report.warnings.append(f"输入不是标准 Markdown，已根据编号/标题样式推断 {inferred} 个标题。")

    return LoadedSpec(markdown_text=ensure_single_trailing_newline(markdown_text), report=report)


def normalize_line_endings(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    return text.lstrip("\ufeff")


def count_markdown_headings(text: str) -> int:
    return sum(1 for line in text.splitlines() if re.match(r"^#{1,6}\s+\S+", line.strip()))


def infer_markdown_headings(text: str) -> tuple[str, int]:
    lines = text.splitlines()
    output: list[str] = []
    inferred = 0
    first_content_seen = False

    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        next_line = lines[index + 1].strip() if index + 1 < len(lines) else ""

        if not stripped:
            output.append("")
            index += 1
            continue

        if re.fullmatch(r"<!--\s*page\s+\d+\s*-->", stripped, flags=re.IGNORECASE):
            output.append(stripped)
            index += 1
            continue

        underline_level = underline_heading_level(next_line)
        if underline_level and is_reasonable_heading(stripped):
            output.append(f"{'#' * underline_level} {stripped}")
            inferred += 1
            first_content_seen = True
            index += 2
            continue

        heading = infer_heading_line(stripped)
        if heading:
            level, title = heading
            output.append(f"{'#' * level} {title}")
            inferred += 1
            first_content_seen = True
            index += 1
            continue

        if not first_content_seen and is_document_title(stripped):
            output.append(f"# {stripped}")
            inferred += 1
            first_content_seen = True
        else:
            output.append(line)
            first_content_seen = True
        index += 1

    return "\n".join(output), inferred


def infer_heading_line(line: str) -> tuple[int, str] | None:
    arabic = re.match(r"^(\d+(?:\.\d+){0,5})[.)、．]?\s+(.+)$", line)
    if arabic and is_reasonable_heading(arabic.group(2)):
        level = min(arabic.group(1).count(".") + 1, 6)
        return level, f"{arabic.group(1)} {arabic.group(2).strip()}"

    chinese_chapter = re.match(r"^(第[一二三四五六七八九十百千万0-9]+[章节])\s*[:：、.-]?\s*(.+)$", line)
    if chinese_chapter and is_reasonable_heading(chinese_chapter.group(2)):
        return 1, f"{chinese_chapter.group(1)} {chinese_chapter.group(2).strip()}"

    chinese_number = re.match(r"^([一二三四五六七八九十]+)[、.．]\s*(.+)$", line)
    if chinese_number and is_reasonable_heading(chinese_number.group(2)):
        return 1, f"{chinese_number.group(1)}、{chinese_number.group(2).strip()}"

    all_caps = re.match(r"^[A-Z][A-Z0-9_ /\-()]{3,}$", line)
    if all_caps and is_reasonable_heading(line):
        return 2, line

    return None


def underline_heading_level(line: str) -> int | None:
    if re.fullmatch(r"={3,}", line):
        return 1
    if re.fullmatch(r"-{3,}", line):
        return 2
    return None


def is_document_title(line: str) -> bool:
    if len(line) > 100:
        return False
    if line.startswith(("-", "*", "|")):
        return False
    return bool(re.search(r"(spec|specification|module|模块|规范|设计|接口)", line, re.IGNORECASE))


def is_reasonable_heading(text: str) -> bool:
    text = text.strip()
    if not text or len(text) > 100:
        return False
    if text.endswith(("。", "；", ";")):
        return False
    if text.startswith(("-", "*", "|", "`")):
        return False
    return True


def markdown_table(rows: list[list[str]]) -> list[str]:
    if not rows:
        return []

    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    body = normalized[1:] or [[""] * width]

    lines = [
        "",
        "| " + " | ".join(escape_table_cell(cell) for cell in header) + " |",
        "| " + " | ".join("---" for _ in range(width)) + " |",
    ]
    for row in body:
        lines.append("| " + " | ".join(escape_table_cell(cell) for cell in row) + " |")
    lines.append("")
    return lines


def escape_table_cell(value: str) -> str:
    return value.replace("|", "\\|")


def ensure_single_trailing_newline(text: str) -> str:
    return text.strip() + "\n"
